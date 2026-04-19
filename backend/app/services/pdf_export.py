from __future__ import annotations

import asyncio
import html
import io
import logging
import re
from dataclasses import dataclass, field
from html.parser import HTMLParser
from typing import Optional, Union

import markdown
from app.services.resume_render import build_render_document

logger = logging.getLogger(__name__)

EXPORT_TIMEOUT_SECONDS = 20
PAGE_TARGETS = {
    "1_page": 1,
    "2_page": 2,
    "3_page": 3,
}
SECTION_HEADING_RE = re.compile(r"^##\s+(.+?)\s*$")
TOP_HEADING_RE = re.compile(r"^#\s+(.+?)\s*$")
SUBHEADING_RE = re.compile(r"^###\s+(.+?)\s*$")
PIPE_ROW_RE = re.compile(r"^(.+?)\s*\|\s*(.+)$")
BULLET_RE = re.compile(r"^\s*[-*+]\s+(.*)$")
EMAIL_RE = re.compile(r"\b[\w.+-]+@[\w.-]+\.[A-Za-z]{2,}\b")
PHONE_RE = re.compile(r"(?:\+\d[\d\s().-]{6,}|\(?\d{3}\)?[\s.-]?\d{3}[\s.-]?\d{4})")
LINKEDIN_RE = re.compile(r"linkedin\.com/|(?:^|[\s|])(?:in|pub|company)/", re.I)
MONTH_YEAR_RE = re.compile(
    r"\b(?:jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|"
    r"jul(?:y)?|aug(?:ust)?|sep(?:t(?:ember)?)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?)\s+\d{4}\b",
    re.I,
)
YEAR_RE = re.compile(r"\b(?:19|20)\d{2}\b")
PRESENT_RE = re.compile(r"\b(?:present|current)\b", re.I)
DATE_RANGE_SEPARATOR_RE = re.compile(r"\bto\b|[-/–—]", re.I)
PROFESSIONAL_EXPERIENCE_HEADING = "professional experience"
ONE_PAGE_VALIDATION_ROOMIER_STEPS = 6
MIN_READABLE_BODY_FONT_SIZE = 9.4
MIN_READABLE_LINE_HEIGHT = 1.1
SINGLE_LIST_WRAPPER_RE = re.compile(r"^<(?:ul|ol)>\s*<li>(.*)</li>\s*</(?:ul|ol)>$", re.S)


@dataclass(frozen=True)
class LayoutPreset:
    body_font_size: float
    line_height: float
    page_margin: float
    spacing_scale: float = 1.0
    section_spacing_scale: float = 1.0

    @property
    def name_font_size(self) -> float:
        return round(self.body_font_size * 1.58, 2)

    @property
    def contact_font_size(self) -> float:
        return round(self.body_font_size * 0.83, 2)

    @property
    def section_heading_size(self) -> float:
        return round(self.body_font_size * 0.95, 2)

    @property
    def section_margin_top(self) -> float:
        return round(self.body_font_size * 0.72 * self.spacing_scale * self.section_spacing_scale, 2)

    @property
    def section_margin_bottom(self) -> float:
        return round(max(4.0, self.body_font_size * 0.32 * self.spacing_scale * self.section_spacing_scale), 2)

    @property
    def paragraph_margin(self) -> float:
        return round(max(1.55, self.body_font_size * 0.18 * self.spacing_scale), 2)

    @property
    def split_row_gap(self) -> float:
        return round(self.body_font_size * 0.52 * self.spacing_scale, 2)

    @property
    def header_margin_bottom(self) -> float:
        return self.contact_to_first_section_margin

    @property
    def contact_to_first_section_margin(self) -> float:
        return round(max(10.0, self.body_font_size * 0.94 * self.spacing_scale), 2)

    @property
    def section_header_content_gap(self) -> float:
        return round(max(4.0, self.body_font_size * 0.42 * self.spacing_scale * self.section_spacing_scale), 2)

    @property
    def subheading_margin_bottom(self) -> float:
        return self.subheading_content_gap

    @property
    def subheading_content_gap(self) -> float:
        return round(max(2.1, self.paragraph_margin * 1.15), 2)

    @property
    def split_group_margin(self) -> float:
        return round(max(0.8, self.paragraph_margin * 0.92), 2)

    @property
    def list_item_margin_bottom(self) -> float:
        return round(max(0.24, self.paragraph_margin * 0.34), 2)

    @property
    def bullet_indent(self) -> float:
        return round(max(10.5, self.body_font_size * 0.96), 2)


@dataclass(frozen=True)
class DocxLayoutPreset:
    body_font_size: float
    line_spacing: float
    page_margin: float
    paragraph_spacing: float
    section_spacing_before: float
    section_spacing_after: float
    header_spacing_after: float
    bullet_indent: float
    split_row_spacing: float

    @property
    def name_font_size(self) -> float:
        return round(self.body_font_size * 1.55, 2)

    @property
    def contact_font_size(self) -> float:
        return round(self.body_font_size * 0.84, 2)

    @property
    def section_heading_size(self) -> float:
        return round(self.body_font_size * 0.95, 2)


@dataclass(frozen=True)
class ExportHeader:
    name: str
    contact_line: str


@dataclass(frozen=True)
class ExportParagraph:
    html_fragment: str


@dataclass(frozen=True)
class ExportSubheading:
    text: str


@dataclass(frozen=True)
class ExportBulletList:
    items: list[str]


@dataclass(frozen=True)
class ExportSplitRow:
    left_html: str
    right_html: str
    emphasize_left: bool = False


@dataclass(frozen=True)
class ExportSplitGroup:
    rows: list[ExportSplitRow]


@dataclass(frozen=True)
class ExportStructuredEntry:
    primary_left_html: str
    primary_right_html: Optional[str]
    secondary_left_html: str
    secondary_right_html: Optional[str]
    bullets: list[str] = field(default_factory=list)
    variant: str = "experience"


ExportBlock = Union[ExportParagraph, ExportSubheading, ExportBulletList, ExportSplitGroup, ExportStructuredEntry]


@dataclass(frozen=True)
class ExportSection:
    heading: str
    blocks: list[ExportBlock]


@dataclass(frozen=True)
class ExportDocument:
    header: Optional[ExportHeader]
    intro_blocks: list[ExportBlock]
    sections: list[ExportSection]
    density_metrics: dict[str, float | int | bool | str]
    normalized_markdown: str


LAYOUT_PRESETS = [
    LayoutPreset(body_font_size=11.2, line_height=1.28, page_margin=0.60, spacing_scale=1.22, section_spacing_scale=1.24),
    LayoutPreset(body_font_size=11.0, line_height=1.24, page_margin=0.56, spacing_scale=1.10, section_spacing_scale=1.14),
    LayoutPreset(body_font_size=10.8, line_height=1.22, page_margin=0.54, spacing_scale=1.04, section_spacing_scale=1.10),
    LayoutPreset(body_font_size=10.5, line_height=1.20, page_margin=0.50, spacing_scale=0.98, section_spacing_scale=1.06),
    LayoutPreset(body_font_size=10.2, line_height=1.18, page_margin=0.48, spacing_scale=0.94, section_spacing_scale=1.02),
    LayoutPreset(body_font_size=9.9, line_height=1.16, page_margin=0.46, spacing_scale=0.88, section_spacing_scale=0.98),
    LayoutPreset(body_font_size=9.6, line_height=1.14, page_margin=0.44, spacing_scale=0.82, section_spacing_scale=0.94),
    LayoutPreset(body_font_size=9.5, line_height=1.12, page_margin=0.42, spacing_scale=0.76, section_spacing_scale=0.90),
    LayoutPreset(body_font_size=9.4, line_height=1.10, page_margin=0.40, spacing_scale=0.72, section_spacing_scale=0.86),
]

DOCX_LAYOUT_PRESETS = {
    "1_page": DocxLayoutPreset(
        body_font_size=10.1,
        line_spacing=1.06,
        page_margin=0.55,
        paragraph_spacing=3.6,
        section_spacing_before=10.8,
        section_spacing_after=6.4,
        header_spacing_after=10.4,
        bullet_indent=18.0,
        split_row_spacing=3.4,
    ),
    "2_page": DocxLayoutPreset(
        body_font_size=10.5,
        line_spacing=1.10,
        page_margin=0.65,
        paragraph_spacing=4.0,
        section_spacing_before=11.2,
        section_spacing_after=6.8,
        header_spacing_after=10.8,
        bullet_indent=18.0,
        split_row_spacing=3.6,
    ),
    "3_page": DocxLayoutPreset(
        body_font_size=10.8,
        line_spacing=1.12,
        page_margin=0.72,
        paragraph_spacing=4.2,
        section_spacing_before=11.6,
        section_spacing_after=7.0,
        header_spacing_after=11.2,
        bullet_indent=18.0,
        split_row_spacing=3.8,
    ),
}


def _clean_personal_value(value: object) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return value.strip()
    return str(value).strip()


def _format_linkedin_value(value: object) -> str:
    linkedin = _clean_personal_value(value)
    if not linkedin:
        return ""

    normalized = re.sub(r"^https?://", "", linkedin, flags=re.I)
    normalized = re.sub(r"^www\.", "", normalized, flags=re.I).rstrip("/")
    match = re.search(r"linkedin\.com/(in|pub|company)/(.+)", normalized, re.I)
    if match:
        return f"{match.group(1).lower()}/{match.group(2).strip('/')}"
    return normalized


def _build_contact_parts(personal_info: Optional[dict]) -> list[str]:
    if not personal_info:
        return []

    parts: list[str] = []
    for key in ("email", "phone", "address"):
        value = _clean_personal_value(personal_info.get(key))
        if value:
            parts.append(value)

    linkedin = _format_linkedin_value(personal_info.get("linkedin_url"))
    if linkedin:
        parts.append(linkedin)
    return parts


def _build_header_lines(personal_info: Optional[dict]) -> list[str]:
    if not personal_info:
        return []

    lines: list[str] = []
    name = _clean_personal_value(personal_info.get("name"))
    if name:
        lines.append(f"# {name}")

    contact_parts = _build_contact_parts(personal_info)
    if contact_parts:
        lines.append(" | ".join(contact_parts))

    return lines


def _looks_like_contact_line(line: str) -> bool:
    stripped = line.strip()
    if not stripped:
        return False
    if "|" in stripped:
        return True
    return bool(
        EMAIL_RE.search(stripped)
        or PHONE_RE.search(stripped)
        or LINKEDIN_RE.search(stripped)
    )


def _join_markdown_parts(header_lines: list[str], body_lines: list[str]) -> str:
    body = "\n".join(body_lines).strip("\n")
    parts: list[str] = []
    if header_lines:
        parts.append("\n".join(header_lines))
    if body:
        parts.append(body)
    if not parts:
        return ""
    return "\n\n".join(parts).rstrip("\n") + "\n"


def _normalize_markdown_for_export(markdown_content: str, personal_info: Optional[dict] = None) -> str:
    stripped_markdown = markdown_content.strip("\n")
    if not stripped_markdown:
        return ""

    lines = stripped_markdown.splitlines()
    first_section_index = next(
        (index for index, line in enumerate(lines) if SECTION_HEADING_RE.match(line.strip())),
        len(lines),
    )
    preamble_lines = lines[:first_section_index]
    body_lines = lines[first_section_index:]

    if not personal_info:
        return stripped_markdown + "\n"

    replacement_header = _build_header_lines(personal_info)
    if not replacement_header:
        return stripped_markdown + "\n"

    nonblank_preamble = [line.strip() for line in preamble_lines if line.strip()]
    profile_name = _clean_personal_value(personal_info.get("name"))

    if not nonblank_preamble:
        return _join_markdown_parts(replacement_header, body_lines)

    first_line = nonblank_preamble[0]
    has_top_heading = bool(TOP_HEADING_RE.match(first_line))
    title = TOP_HEADING_RE.match(first_line).group(1).strip() if has_top_heading else ""
    remaining_lines = nonblank_preamble[1:] if has_top_heading else nonblank_preamble
    plain_remaining_lines = nonblank_preamble[1:]
    all_remaining_contactish = all(_looks_like_contact_line(line) for line in remaining_lines)

    should_replace = False
    if title in {"(Name)", "Resume", "CV"}:
        should_replace = True
    elif any("invite-only@" in line.lower() for line in nonblank_preamble):
        should_replace = True
    elif has_top_heading and profile_name and title == profile_name and all_remaining_contactish:
        should_replace = True
    elif not has_top_heading and profile_name and first_line == profile_name and all(_looks_like_contact_line(line) for line in plain_remaining_lines):
        should_replace = True
    elif not has_top_heading and all(_looks_like_contact_line(line) for line in nonblank_preamble):
        should_replace = True

    if should_replace:
        return _join_markdown_parts(replacement_header, body_lines)

    if not has_top_heading and body_lines:
        return _join_markdown_parts(replacement_header, preamble_lines + body_lines)

    return stripped_markdown + "\n"


def _unwrap_single_paragraph(html_fragment: str) -> str:
    stripped = html_fragment.strip()
    if stripped.startswith("<p>") and stripped.endswith("</p>") and stripped.count("<p>") == 1:
        return stripped[3:-4]
    return stripped


def _render_inline_markdown(text: str) -> str:
    rendered = markdown.markdown(text, extensions=["extra", "sane_lists"])
    return _unwrap_single_paragraph(rendered)


def _render_markdown_block(text: str) -> str:
    return markdown.markdown(text, extensions=["extra", "sane_lists"]).strip()


def _is_professional_experience_section(section_heading: Optional[str]) -> bool:
    return bool(section_heading and section_heading.strip().lower() == PROFESSIONAL_EXPERIENCE_HEADING)


def _looks_like_experience_date_range(value: str) -> bool:
    normalized = re.sub(r"\s+", " ", value).strip()
    if not normalized:
        return False
    has_month_year = bool(MONTH_YEAR_RE.search(normalized))
    has_year = bool(YEAR_RE.search(normalized))
    has_present = bool(PRESENT_RE.search(normalized))
    has_range_separator = bool(DATE_RANGE_SEPARATOR_RE.search(normalized))
    return (has_month_year or has_year) and (has_range_separator or has_present)


def _render_split_row(left: str, right: str, *, emphasize_left: bool = False) -> str:
    row_class = "split-row split-row-role-title" if emphasize_left else "split-row"
    left_class = "split-left split-left-strong" if emphasize_left else "split-left"
    return (
        f"<div class='{row_class}'>"
        f"<span class='{left_class}'>{_render_inline_markdown(left)}</span>"
        f"<span class='split-right'>{_render_inline_markdown(right)}</span>"
        "</div>"
    )


def _render_list_item_content(text: str) -> str:
    rendered = _render_inline_markdown(text)
    stripped = rendered.strip()
    match = SINGLE_LIST_WRAPPER_RE.fullmatch(stripped)
    if match:
        return match.group(1).strip()
    return stripped


def _calculate_content_density_metrics(markdown_content: str) -> dict[str, float | int | bool | str]:
    lines = [line.strip() for line in markdown_content.strip().splitlines() if line.strip()]
    total_lines = len(lines)
    bullet_count = sum(1 for line in lines if BULLET_RE.match(line))
    section_count = sum(1 for line in lines if SECTION_HEADING_RE.match(line))
    content_line_count = sum(
        1
        for line in lines
        if not SECTION_HEADING_RE.match(line) and not TOP_HEADING_RE.match(line)
    )
    bullets_per_section = bullet_count / max(1, section_count)
    lines_per_section = content_line_count / max(1, section_count)
    is_dense = total_lines >= 30 or bullets_per_section >= 4.5 or lines_per_section >= 8.0
    is_sparse = total_lines <= 18 and bullet_count <= 6 and lines_per_section <= 5.0
    density_label = "dense" if is_dense else "sparse" if is_sparse else "balanced"

    return {
        "total_lines": total_lines,
        "bullet_count": bullet_count,
        "section_count": section_count,
        "content_line_count": content_line_count,
        "bullets_per_section": round(bullets_per_section, 2),
        "lines_per_section": round(lines_per_section, 2),
        "is_dense": is_dense,
        "is_sparse": is_sparse,
        "density_label": density_label,
    }


def _parse_content_blocks(lines: list[str], *, section_heading: Optional[str] = None) -> list[ExportBlock]:
    blocks: list[ExportBlock] = []
    index = 0
    in_professional_experience = _is_professional_experience_section(section_heading)

    while index < len(lines):
        stripped = lines[index].strip()
        if not stripped:
            index += 1
            continue

        subheading_match = SUBHEADING_RE.match(stripped)
        if subheading_match:
            blocks.append(ExportSubheading(text=subheading_match.group(1).strip()))
            index += 1
            continue

        if not BULLET_RE.match(stripped):
            pipe_match = PIPE_ROW_RE.match(stripped)
            if pipe_match:
                rows: list[ExportSplitRow] = []
                while index < len(lines):
                    current = lines[index].strip()
                    current_match = PIPE_ROW_RE.match(current)
                    if not current or BULLET_RE.match(current) or current_match is None:
                        break
                    right_column = current_match.group(2).strip()
                    rows.append(
                        ExportSplitRow(
                            left_html=_render_inline_markdown(current_match.group(1).strip()),
                            right_html=_render_inline_markdown(right_column),
                            emphasize_left=(
                                in_professional_experience
                                and _looks_like_experience_date_range(right_column)
                            ),
                        )
                    )
                    index += 1
                if rows:
                    blocks.append(ExportSplitGroup(rows=rows))
                continue

        bullet_match = BULLET_RE.match(stripped)
        if bullet_match:
            items: list[str] = []
            while index < len(lines):
                current = lines[index].strip()
                current_match = BULLET_RE.match(current)
                if current_match is None:
                    break
                items.append(_render_list_item_content(current_match.group(1).strip()))
                index += 1
            blocks.append(ExportBulletList(items=items))
            continue

        paragraph_lines = [stripped]
        index += 1
        while index < len(lines):
            current = lines[index].strip()
            if (
                not current
                or SUBHEADING_RE.match(current)
                or BULLET_RE.match(current)
                or PIPE_ROW_RE.match(current)
            ):
                break
            paragraph_lines.append(current)
            index += 1

        blocks.append(ExportParagraph(html_fragment=_render_markdown_block("\n".join(paragraph_lines))))

    return blocks


def _render_content_blocks(lines: list[str], *, section_heading: Optional[str] = None) -> str:
    blocks = _parse_content_blocks(lines, section_heading=section_heading)
    return _render_html_blocks(blocks)


def _build_export_document(markdown_content: str, personal_info: Optional[dict] = None) -> ExportDocument:
    normalized_markdown = _normalize_markdown_for_export(markdown_content, personal_info)
    render_result = build_render_document(normalized_markdown)
    if render_result.document is None:
        raise ValueError(render_result.error or "Resume render model could not be built for export.")

    render_document = render_result.document
    density_metrics = _calculate_content_density_metrics(render_document.normalized_markdown)
    header = None
    if render_document.header and (render_document.header.name or render_document.header.contact_line):
        header = ExportHeader(
            name=str(render_document.header.name or ""),
            contact_line=str(render_document.header.contact_line or ""),
        )

    intro_blocks = _parse_content_blocks(list(render_document.header.extra_lines)) if render_document.header else []
    sections: list[ExportSection] = []
    for section in render_document.sections:
        if section.kind in {"professional_experience", "education"}:
            structured_blocks: list[ExportBlock] = []
            for entry in section.entries:
                structured_blocks.append(
                    ExportStructuredEntry(
                        primary_left_html=_render_inline_markdown(entry.row1_left),
                        primary_right_html=(
                            _render_inline_markdown(entry.row1_right) if entry.row1_right else None
                        ),
                        secondary_left_html=_render_inline_markdown(entry.row2_left),
                        secondary_right_html=(
                            _render_inline_markdown(entry.row2_right) if entry.row2_right else None
                        ),
                        bullets=[_render_list_item_content(bullet) for bullet in entry.bullets],
                        variant=section.kind,
                    )
                )
            sections.append(ExportSection(heading=section.heading, blocks=structured_blocks))
            continue
        section_lines = section.markdown_body.splitlines() if section.markdown_body else []
        sections.append(
            ExportSection(
                heading=section.heading,
                blocks=_parse_content_blocks(section_lines, section_heading=section.heading),
            )
        )

    return ExportDocument(
        header=header,
        intro_blocks=intro_blocks,
        sections=sections,
        density_metrics=density_metrics,
        normalized_markdown=render_document.normalized_markdown,
    )


def _render_html_blocks(blocks: list[ExportBlock]) -> str:
    html_blocks: list[str] = []

    def render_split_group(group: ExportSplitGroup) -> str:
        rows = []
        for row in group.rows:
            row_class = "split-row split-row-role-title" if row.emphasize_left else "split-row"
            left_class = "split-left split-left-strong" if row.emphasize_left else "split-left"
            rows.append(
                f"<div class='{row_class}'>"
                f"<span class='{left_class}'>{row.left_html}</span>"
                f"<span class='split-right'>{row.right_html}</span>"
                "</div>"
            )
        return f"<div class='split-group'>{''.join(rows)}</div>"

    def render_structured_entry(entry: ExportStructuredEntry) -> str:
        bullet_html = ""
        if entry.bullets:
            items = "".join(f"<li>{item}</li>" for item in entry.bullets)
            bullet_html = f"<ul>{items}</ul>"
        return (
            f"<div class='structured-entry structured-entry-{entry.variant}'>"
            "<div class='structured-row structured-row-primary'>"
            f"<span class='structured-left structured-left-primary'>{entry.primary_left_html}</span>"
            f"<span class='structured-right structured-right-primary'>{entry.primary_right_html or ''}</span>"
            "</div>"
            "<div class='structured-row structured-row-secondary'>"
            f"<span class='structured-left structured-left-secondary'>{entry.secondary_left_html}</span>"
            f"<span class='structured-right structured-right-secondary'>{entry.secondary_right_html or ''}</span>"
            "</div>"
            f"{bullet_html}"
            "</div>"
        )

    for block in blocks:
        if isinstance(block, ExportParagraph):
            html_blocks.append(block.html_fragment)
        elif isinstance(block, ExportSubheading):
            html_blocks.append(f"<h3>{html.escape(block.text)}</h3>")
        elif isinstance(block, ExportBulletList):
            items = "".join(f"<li>{item}</li>" for item in block.items)
            html_blocks.append(f"<ul>{items}</ul>")
        elif isinstance(block, ExportSplitGroup):
            html_blocks.append(render_split_group(block))
        elif isinstance(block, ExportStructuredEntry):
            html_blocks.append(render_structured_entry(block))
    return "".join(html_blocks)


def _build_html(
    document_or_markdown: Union[ExportDocument, str],
    preset: LayoutPreset,
    *,
    preset_index: int = 0,
) -> str:
    document = (
        document_or_markdown
        if isinstance(document_or_markdown, ExportDocument)
        else _build_export_document(document_or_markdown)
    )
    density_metrics = document.density_metrics
    density_label = str(density_metrics["density_label"])
    spacing_factor = 0.92 if density_metrics["is_dense"] else 1.08 if density_metrics["is_sparse"] else 1.0
    section_spacing_factor = 0.9 if density_metrics["is_dense"] else 1.1 if density_metrics["is_sparse"] else 1.0
    major_section_gap = round(
        max(
            preset.body_font_size * 0.6 * preset.section_spacing_scale,
            preset.section_margin_top * section_spacing_factor,
        ),
        2,
    )
    header_gap = round(max(major_section_gap, preset.contact_to_first_section_margin * section_spacing_factor), 2)
    section_margin_top = major_section_gap
    section_heading_gap = round(max(4.0, preset.section_header_content_gap * section_spacing_factor), 2)
    subheading_gap = round(max(2.1, preset.subheading_content_gap * spacing_factor), 2)
    paragraph_margin = round(max(1.4, preset.paragraph_margin * spacing_factor), 2)
    split_group_margin = round(max(0.8, preset.split_group_margin * spacing_factor), 2)
    list_item_margin = round(max(0.24, preset.list_item_margin_bottom * spacing_factor), 2)
    bullet_indent = round(max(10.5, preset.bullet_indent), 2)
    bullet_padding = round(max(3.5, bullet_indent * 0.36), 2)
    content_font_size = round(preset.body_font_size * 0.86, 2)
    subheader_font_size = round(
        min(
            preset.section_heading_size - 0.18,
            max(content_font_size + 0.32, preset.body_font_size * 0.90),
        ),
        2,
    )
    structured_entry_gap = round(max(split_group_margin * 3.0, paragraph_margin * 3.2), 2)

    intro_html = _render_html_blocks(document.intro_blocks)
    sections_html = "".join(
        (
            "<section class='resume-section'>"
            f"<h2>{html.escape(section.heading)}</h2>"
            f"{_render_html_blocks(section.blocks)}"
            "</section>"
        )
        for section in document.sections
    )
    header_html = ""
    if document.header and (document.header.name or document.header.contact_line):
        contact_html = (
            f"<p class='resume-contact'>{html.escape(document.header.contact_line)}</p>"
            if document.header.contact_line
            else ""
        )
        header_html = (
            "<header class='resume-header'>"
            f"<h1>{html.escape(document.header.name)}</h1>"
            f"{contact_html}"
            "</header>"
        )

    return f"""<!DOCTYPE html>
<html>
<head>
  <meta charset="utf-8">
  <style>
    @page {{
      size: Letter;
      margin: {preset.page_margin}in;
    }}
    body {{
      margin: 0;
      font-family: 'Georgia', 'Times New Roman', serif;
      font-size: {preset.body_font_size}pt;
      line-height: {preset.line_height};
      color: #111111;
    }}
    .resume-root {{
      width: 100%;
    }}
    .resume-header {{
      text-align: center;
      margin-bottom: {header_gap}pt;
    }}
    .resume-header h1 {{
      font-size: {preset.name_font_size}pt;
      margin: 0 0 0.5pt 0;
      font-weight: 700;
      line-height: 1;
    }}
    .resume-contact {{
      font-size: {preset.contact_font_size}pt;
      margin: 0;
      line-height: 1.08;
    }}
    .resume-section {{
      margin-top: {section_margin_top}pt;
    }}
    .resume-section:first-of-type {{
      margin-top: 0;
    }}
    .resume-section h2 {{
      margin: 0 0 {section_heading_gap}pt 0;
      padding-bottom: 0.5pt;
      font-size: {preset.section_heading_size}pt;
      font-weight: 700;
      letter-spacing: 0.015em;
      text-transform: uppercase;
      line-height: 1;
      border-bottom: 0.8pt solid #111111;
    }}
    h3 {{
      margin: 0 0 {subheading_gap}pt 0;
      font-size: {subheader_font_size}pt;
      font-weight: 700;
      line-height: 1.05;
    }}
    p {{
      margin: 0 0 {paragraph_margin}pt 0;
      font-size: {content_font_size}pt;
    }}
    ul {{
      margin: 0 0 {paragraph_margin}pt {bullet_indent}pt;
      padding-left: {bullet_padding}pt;
      font-size: {content_font_size}pt;
    }}
    li {{
      margin: 0 0 {list_item_margin}pt 0;
      padding-left: 0;
      font-size: {content_font_size}pt;
    }}
    .split-group {{
      margin: 0 0 {split_group_margin}pt 0;
    }}
    .structured-entry {{
      margin: 0 0 {structured_entry_gap}pt 0;
    }}
    .structured-row {{
      display: flex;
      justify-content: space-between;
      align-items: baseline;
      gap: {preset.split_row_gap}pt;
      width: 100%;
      margin: 0;
    }}
    .structured-row + .structured-row {{
      margin-top: {round(max(1.6, paragraph_margin * 0.45), 2)}pt;
    }}
    .structured-left {{
      flex: 1 1 auto;
      min-width: 0;
    }}
    .structured-right {{
      flex: 0 0 auto;
      text-align: right;
      white-space: nowrap;
      font-style: italic;
    }}
    .structured-left-primary,
    .structured-right-primary {{
      font-size: {subheader_font_size}pt;
      line-height: 1.02;
    }}
    .structured-left-secondary,
    .structured-right-secondary {{
      font-size: {subheader_font_size}pt;
      line-height: 1.02;
    }}
    .structured-left-primary {{
      font-weight: 700;
    }}
    .structured-left-secondary {{
      font-style: italic;
    }}
    .structured-entry ul {{
      margin-top: {round(max(1.8, paragraph_margin * 0.5), 2)}pt;
    }}
    .split-row {{
      display: flex;
      justify-content: space-between;
      align-items: baseline;
      gap: {preset.split_row_gap}pt;
      width: 100%;
      margin: 0;
    }}
    .split-left {{
      flex: 1 1 auto;
      min-width: 0;
    }}
    .split-right {{
      flex: 0 0 auto;
      text-align: right;
      white-space: nowrap;
    }}
    .split-left-strong {{
      font-weight: 700;
    }}
    .split-group + ul {{
      margin-top: 0;
    }}
    .split-group + .split-group {{
      margin-top: {preset.split_group_margin}pt;
    }}
    .resume-section > :last-child {{
      margin-bottom: 0;
    }}
    strong {{
      font-weight: 700;
    }}
    em {{
      font-style: italic;
    }}
  </style>
</head>
<body data-preset="{preset_index}" data-density="{density_label}">
  <main class="resume-root resume-root-{density_label}">
    {header_html}
    {intro_html}
    {sections_html}
  </main>
</body>
</html>"""


def _render_html_to_pdf(html_content: str) -> tuple[bytes, int]:
    import weasyprint  # noqa: WPS433

    document = weasyprint.HTML(string=html_content).render()
    return document.write_pdf(), len(document.pages)


def _is_readable_preset(preset: LayoutPreset) -> bool:
    return preset.body_font_size >= MIN_READABLE_BODY_FONT_SIZE and preset.line_height >= MIN_READABLE_LINE_HEIGHT


def _build_roomier_one_page_variant(preset: LayoutPreset) -> LayoutPreset:
    return LayoutPreset(
        body_font_size=round(min(12.8, preset.body_font_size + 0.22), 2),
        line_height=round(min(1.32, preset.line_height + 0.02), 2),
        page_margin=preset.page_margin,
        spacing_scale=round(min(1.35, preset.spacing_scale + 0.07), 2),
        section_spacing_scale=round(min(1.75, preset.section_spacing_scale + 0.08), 2),
    )


def _build_section_relief_one_page_variant(preset: LayoutPreset) -> LayoutPreset:
    return LayoutPreset(
        body_font_size=preset.body_font_size,
        line_height=preset.line_height,
        page_margin=preset.page_margin,
        spacing_scale=preset.spacing_scale,
        section_spacing_scale=round(min(1.9, preset.section_spacing_scale + 0.14), 2),
    )


def _validate_roomier_one_page_fit(
    document: ExportDocument,
    *,
    base_pdf: bytes,
    base_preset: LayoutPreset,
    base_index: int,
) -> bytes:
    best_pdf = base_pdf
    best_preset = base_preset

    for _step in range(ONE_PAGE_VALIDATION_ROOMIER_STEPS):
        improved = False
        for build_candidate in (
            _build_section_relief_one_page_variant,
            _build_roomier_one_page_variant,
        ):
            candidate_preset = build_candidate(best_preset)
            if candidate_preset == best_preset:
                continue
            html_content = _build_html(document, candidate_preset, preset_index=base_index)
            pdf_bytes, page_count = _render_html_to_pdf(html_content)
            if page_count > 1:
                continue
            best_preset = candidate_preset
            best_pdf = pdf_bytes
            improved = True
        if not improved:
            break

    return best_pdf


def _generate_pdf_with_autofit_sync(
    markdown_content: str,
    personal_info: Optional[dict] = None,
    page_length: Optional[str] = None,
) -> bytes:
    document = _build_export_document(markdown_content, personal_info)
    target_pages = PAGE_TARGETS.get(str(page_length or "1_page"), 1)

    last_pdf = b""
    for preset_index, preset in enumerate(LAYOUT_PRESETS):
        if not _is_readable_preset(preset):
            continue
        html_content = _build_html(document, preset, preset_index=preset_index)
        pdf_bytes, page_count = _render_html_to_pdf(html_content)
        last_pdf = pdf_bytes
        if page_count <= target_pages:
            if target_pages == 1:
                return _validate_roomier_one_page_fit(
                    document,
                    base_pdf=pdf_bytes,
                    base_preset=preset,
                    base_index=preset_index,
                )
            return pdf_bytes

    return last_pdf


class _InlineHTMLToDocxParser(HTMLParser):
    def __init__(self, paragraph, *, default_bold: bool = False, default_italic: bool = False) -> None:
        super().__init__()
        self.paragraph = paragraph
        self.bold_depth = 1 if default_bold else 0
        self.italic_depth = 1 if default_italic else 0

    def handle_starttag(self, tag: str, attrs) -> None:  # noqa: ANN001
        if tag in {"strong", "b"}:
            self.bold_depth += 1
        elif tag in {"em", "i"}:
            self.italic_depth += 1

    def handle_endtag(self, tag: str) -> None:
        if tag in {"strong", "b"} and self.bold_depth > 0:
            self.bold_depth -= 1
        elif tag in {"em", "i"} and self.italic_depth > 0:
            self.italic_depth -= 1

    def handle_data(self, data: str) -> None:
        if not data:
            return
        run = self.paragraph.add_run(html.unescape(data))
        run.bold = self.bold_depth > 0
        run.italic = self.italic_depth > 0


def _append_inline_html_runs(paragraph, html_fragment: str, *, bold: bool = False, italic: bool = False) -> None:
    parser = _InlineHTMLToDocxParser(paragraph, default_bold=bold, default_italic=italic)
    parser.feed(html_fragment)
    parser.close()


def _set_paragraph_spacing(paragraph, *, before: float = 0.0, after: float = 0.0, line_spacing: Optional[float] = None) -> None:
    from docx.shared import Pt

    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(before)
    paragraph_format.space_after = Pt(after)
    if line_spacing is not None:
        paragraph_format.line_spacing = line_spacing


def _set_font(run, *, size: float, font_name: str = "Georgia", bold: Optional[bool] = None) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt

    run.font.name = font_name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def _apply_font_to_paragraph_runs(paragraph, *, size: float, font_name: str = "Georgia") -> None:
    for run in paragraph.runs:
        _set_font(run, size=size, font_name=font_name)


def _resolve_docx_layout(page_length: Optional[str], density_metrics: dict[str, float | int | bool | str]) -> DocxLayoutPreset:
    base = DOCX_LAYOUT_PRESETS.get(str(page_length or "1_page"), DOCX_LAYOUT_PRESETS["1_page"])
    if density_metrics["is_dense"]:
        return DocxLayoutPreset(
            body_font_size=max(10.0, round(base.body_font_size - 0.3, 2)),
            line_spacing=max(1.0, round(base.line_spacing - 0.03, 2)),
            page_margin=max(0.5, round(base.page_margin - 0.04, 2)),
            paragraph_spacing=max(2.8, round(base.paragraph_spacing - 0.2, 2)),
            section_spacing_before=max(8.8, round(base.section_spacing_before - 0.4, 2)),
            section_spacing_after=max(5.4, round(base.section_spacing_after - 0.2, 2)),
            header_spacing_after=max(8.8, round(base.header_spacing_after - 0.2, 2)),
            bullet_indent=base.bullet_indent,
            split_row_spacing=max(2.8, round(base.split_row_spacing - 0.1, 2)),
        )
    if density_metrics["is_sparse"]:
        return DocxLayoutPreset(
            body_font_size=round(base.body_font_size + 0.25, 2),
            line_spacing=round(base.line_spacing + 0.03, 2),
            page_margin=round(base.page_margin, 2),
            paragraph_spacing=round(base.paragraph_spacing + 0.4, 2),
            section_spacing_before=round(base.section_spacing_before + 0.8, 2),
            section_spacing_after=round(base.section_spacing_after + 0.5, 2),
            header_spacing_after=round(base.header_spacing_after + 0.8, 2),
            bullet_indent=base.bullet_indent,
            split_row_spacing=round(base.split_row_spacing + 0.3, 2),
        )
    return base


def _render_docx_sync(
    markdown_content: str,
    personal_info: Optional[dict] = None,
    page_length: Optional[str] = None,
) -> bytes:
    from docx import Document  # noqa: WPS433
    from docx.enum.section import WD_SECTION_START  # noqa: WPS433
    from docx.enum.text import WD_TAB_ALIGNMENT  # noqa: WPS433
    from docx.shared import Inches, Pt  # noqa: WPS433

    document_model = _build_export_document(markdown_content, personal_info)
    layout = _resolve_docx_layout(page_length, document_model.density_metrics)
    major_section_gap = max(layout.body_font_size * layout.line_spacing * 0.95, layout.section_spacing_before)
    header_section_gap = max(layout.body_font_size * layout.line_spacing * 0.95, layout.header_spacing_after)

    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(layout.page_margin)
    section.bottom_margin = Inches(layout.page_margin)
    section.left_margin = Inches(layout.page_margin)
    section.right_margin = Inches(layout.page_margin)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Georgia"
    normal_style.font.size = Pt(layout.body_font_size)
    normal_style.paragraph_format.line_spacing = layout.line_spacing
    normal_style.paragraph_format.space_before = Pt(0)
    normal_style.paragraph_format.space_after = Pt(layout.paragraph_spacing)

    if document_model.header and (document_model.header.name or document_model.header.contact_line):
        header_paragraph = doc.add_paragraph()
        header_paragraph.alignment = 1
        _set_paragraph_spacing(header_paragraph, after=0, line_spacing=1.0)
        if document_model.header.name:
            name_run = header_paragraph.add_run(document_model.header.name)
            _set_font(name_run, size=layout.name_font_size, bold=True)

        if document_model.header.contact_line:
            contact_paragraph = doc.add_paragraph()
            contact_paragraph.alignment = 1
            _set_paragraph_spacing(contact_paragraph, after=header_section_gap, line_spacing=1.0)
            contact_run = contact_paragraph.add_run(document_model.header.contact_line)
            _set_font(contact_run, size=layout.contact_font_size)

    def add_blocks(blocks: list[ExportBlock]) -> None:
        usable_width = section.page_width - section.left_margin - section.right_margin
        content_font_size = max(9.0, round(layout.body_font_size - 0.45, 2))
        subheader_font_size = min(layout.section_heading_size - 0.15, round(content_font_size + 0.35, 2))
        entry_spacing_after = layout.paragraph_spacing + 2.8

        for block in blocks:
            if isinstance(block, ExportParagraph):
                paragraph = doc.add_paragraph()
                _set_paragraph_spacing(paragraph, after=layout.paragraph_spacing, line_spacing=layout.line_spacing)
                _append_inline_html_runs(paragraph, block.html_fragment)
                _apply_font_to_paragraph_runs(paragraph, size=content_font_size)
            elif isinstance(block, ExportSubheading):
                paragraph = doc.add_paragraph()
                _set_paragraph_spacing(paragraph, after=layout.paragraph_spacing, line_spacing=1.0)
                run = paragraph.add_run(block.text)
                _set_font(run, size=subheader_font_size, bold=True)
            elif isinstance(block, ExportBulletList):
                last_index = len(block.items) - 1
                for item_index, item in enumerate(block.items):
                    paragraph = doc.add_paragraph(style="List Bullet")
                    paragraph.paragraph_format.left_indent = Pt(layout.bullet_indent)
                    paragraph_after = entry_spacing_after if item_index == last_index else layout.paragraph_spacing
                    _set_paragraph_spacing(paragraph, after=paragraph_after, line_spacing=layout.line_spacing)
                    _append_inline_html_runs(paragraph, item)
                    _apply_font_to_paragraph_runs(paragraph, size=content_font_size)
            elif isinstance(block, ExportSplitGroup):
                last_index = len(block.rows) - 1
                for row_index, row in enumerate(block.rows):
                    paragraph = doc.add_paragraph()
                    paragraph.paragraph_format.tab_stops.add_tab_stop(usable_width, WD_TAB_ALIGNMENT.RIGHT)
                    paragraph_after = layout.split_row_spacing if row_index < last_index else entry_spacing_after
                    _set_paragraph_spacing(paragraph, after=paragraph_after, line_spacing=layout.line_spacing)
                    _append_inline_html_runs(paragraph, row.left_html, bold=row.emphasize_left)
                    paragraph.add_run("\t")
                    _append_inline_html_runs(paragraph, row.right_html)
                    _apply_font_to_paragraph_runs(paragraph, size=subheader_font_size if row.emphasize_left else content_font_size)
            elif isinstance(block, ExportStructuredEntry):
                primary = doc.add_paragraph()
                primary.paragraph_format.tab_stops.add_tab_stop(usable_width, WD_TAB_ALIGNMENT.RIGHT)
                _set_paragraph_spacing(primary, after=max(1.4, layout.split_row_spacing - 0.1), line_spacing=1.0)
                _append_inline_html_runs(primary, block.primary_left_html, bold=True)
                if block.primary_right_html:
                    primary.add_run("\t")
                    _append_inline_html_runs(primary, block.primary_right_html, italic=True)
                _apply_font_to_paragraph_runs(primary, size=subheader_font_size)

                secondary = doc.add_paragraph()
                secondary.paragraph_format.tab_stops.add_tab_stop(usable_width, WD_TAB_ALIGNMENT.RIGHT)
                secondary_after = entry_spacing_after if not block.bullets else max(2.2, layout.paragraph_spacing + 0.4)
                _set_paragraph_spacing(secondary, after=secondary_after, line_spacing=1.0)
                _append_inline_html_runs(secondary, block.secondary_left_html, italic=True)
                if block.secondary_right_html:
                    secondary.add_run("\t")
                    _append_inline_html_runs(secondary, block.secondary_right_html, italic=True)
                _apply_font_to_paragraph_runs(secondary, size=subheader_font_size)

                last_index = len(block.bullets) - 1
                for item_index, item in enumerate(block.bullets):
                    paragraph = doc.add_paragraph(style="List Bullet")
                    paragraph.paragraph_format.left_indent = Pt(layout.bullet_indent)
                    paragraph_after = entry_spacing_after if item_index == last_index else layout.paragraph_spacing
                    _set_paragraph_spacing(paragraph, after=paragraph_after, line_spacing=layout.line_spacing)
                    _append_inline_html_runs(paragraph, item)
                    _apply_font_to_paragraph_runs(paragraph, size=content_font_size)

    add_blocks(document_model.intro_blocks)

    for section_block in document_model.sections:
        heading_paragraph = doc.add_paragraph()
        _set_paragraph_spacing(
            heading_paragraph,
            before=major_section_gap,
            after=layout.section_spacing_after,
            line_spacing=1.0,
        )
        heading_run = heading_paragraph.add_run(section_block.heading.upper())
        _set_font(heading_run, size=layout.section_heading_size, bold=True)
        bottom_border = heading_paragraph._element.get_or_add_pPr()
        # Word border XML is a minimal way to preserve the PDF heading separator without tables/shapes.
        from docx.oxml import OxmlElement  # noqa: WPS433
        from docx.oxml.ns import qn  # noqa: WPS433

        p_borders = OxmlElement("w:pBdr")
        border = OxmlElement("w:bottom")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "6")
        border.set(qn("w:space"), "1")
        border.set(qn("w:color"), "111111")
        p_borders.append(border)
        bottom_border.append(p_borders)
        add_blocks(section_block.blocks)

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


async def generate_pdf(
    markdown_content: str,
    personal_info: Optional[dict] = None,
    page_length: Optional[str] = None,
) -> bytes:
    loop = asyncio.get_running_loop()
    return await asyncio.wait_for(
        loop.run_in_executor(
            None,
            _generate_pdf_with_autofit_sync,
            markdown_content,
            personal_info,
            page_length,
        ),
        timeout=EXPORT_TIMEOUT_SECONDS,
    )


async def generate_docx(
    markdown_content: str,
    personal_info: Optional[dict] = None,
    page_length: Optional[str] = None,
) -> bytes:
    loop = asyncio.get_running_loop()
    return await asyncio.wait_for(
        loop.run_in_executor(
            None,
            _render_docx_sync,
            markdown_content,
            personal_info,
            page_length,
        ),
        timeout=EXPORT_TIMEOUT_SECONDS,
    )
