from __future__ import annotations

import io
import re

from app.services import pdf_export


def _personal_info() -> dict[str, str]:
    return {
        "name": "Alex Example",
        "email": "alex@example.com",
        "phone": "555-0100",
        "address": "Toronto, ON",
        "linkedin_url": "https://linkedin.com/in/alex-example",
    }


def test_normalize_markdown_replaces_legacy_placeholder_header():
    normalized = pdf_export._normalize_markdown_for_export(
        "# (Name)\ninvite-only@example.com\n\n## Summary\nBuilt backend systems.\n",
        _personal_info(),
    )

    assert normalized.startswith(
        "# Alex Example\nalex@example.com | 555-0100 | Toronto, ON | in/alex-example\n\n## Summary\n"
    )
    assert "(Name)" not in normalized
    assert "invite-only@example.com" not in normalized


def test_build_html_renders_one_header_for_existing_assembly_header():
    markdown_content = (
        "# Alex Example\n"
        "alex@example.com | 555-0100 | Toronto, ON | in/alex-example\n\n"
        "## Summary\n"
        "Built backend systems.\n"
    )

    html = pdf_export._build_html(markdown_content, pdf_export.LAYOUT_PRESETS[0])

    assert html.count("<header class=\"resume-header\">") == 0
    assert html.count("<header class='resume-header'>") == 1
    assert html.count("<h1>Alex Example</h1>") == 1
    assert "## Summary" not in html
    assert "Built backend systems." in html


def test_normalize_markdown_replaces_plain_text_header_matching_profile():
    normalized = pdf_export._normalize_markdown_for_export(
        "Alex Example\nalex@example.com | 555-0100\n\n## Summary\nBuilt backend systems.\n",
        _personal_info(),
    )

    assert normalized.startswith(
        "# Alex Example\nalex@example.com | 555-0100 | Toronto, ON | in/alex-example\n\n## Summary\n"
    )
    assert "\n\nAlex Example\nalex@example.com | 555-0100\n" not in normalized


def test_build_html_uses_point_spacing_units():
    preset = pdf_export.LAYOUT_PRESETS[0]
    html = pdf_export._build_html(
        "## Summary\nBuilt backend systems.\n",
        preset,
    )
    major_section_gap = max(
        preset.body_font_size * 0.6 * preset.section_spacing_scale,
        preset.section_margin_top * 1.1,
    )
    header_gap = max(major_section_gap, preset.contact_to_first_section_margin * 1.1)

    assert f"margin-bottom: {header_gap:.2f}pt;" in html
    assert f"margin-top: {major_section_gap:.2f}pt;" in html
    assert f"margin: 0 0 {preset.section_header_content_gap * 1.1:.2f}pt 0;" in html
    assert f"margin: 0 0 {preset.paragraph_margin * 1.08:.2f}pt 0;" in html
    assert f"margin: 0 0 {preset.paragraph_margin * 1.08:.2f}pt {preset.bullet_indent}pt;" in html
    assert "margin-top: 0;" in html
    assert 'data-density="sparse"' in html
    assert f"{preset.section_margin_top}rem" not in html
    assert f"{preset.section_header_content_gap}rem" not in html


def test_render_content_blocks_unwraps_single_item_list_markup_inside_bullets():
    html = pdf_export._render_content_blocks(["- * nested bullet"])

    assert html == "<ul><li>nested bullet</li></ul>"


def test_render_content_blocks_preserves_literal_star_content_inside_bullets():
    html = pdf_export._render_content_blocks(["- *nix operations"])

    assert "*nix operations" in html
    assert "<ul><li><ul>" not in html


def test_calculate_content_density_metrics_flags_dense_and_sparse_documents():
    dense_metrics = pdf_export._calculate_content_density_metrics(
        (
            "## Experience\n"
            "- one\n- two\n- three\n- four\n- five\n"
            "## Projects\n"
            "- six\n- seven\n- eight\n- nine\n- ten\n"
            "## Skills\n"
            "- python\n- sql\n- aws\n- docker\n- linux\n"
        )
    )
    sparse_metrics = pdf_export._calculate_content_density_metrics(
        "## Summary\nBuilt backend systems.\n## Skills\nPython | SQL\n"
    )

    assert dense_metrics["is_dense"] is True
    assert dense_metrics["density_label"] == "dense"
    assert sparse_metrics["is_sparse"] is True
    assert sparse_metrics["density_label"] == "sparse"


def test_layout_presets_keep_readable_minimums():
    assert all(preset.body_font_size >= pdf_export.MIN_READABLE_BODY_FONT_SIZE for preset in pdf_export.LAYOUT_PRESETS)
    assert all(preset.line_height >= pdf_export.MIN_READABLE_LINE_HEIGHT for preset in pdf_export.LAYOUT_PRESETS)


def test_generate_pdf_autofit_retries_until_target_page_count_is_met(monkeypatch):
    seen_presets: list[int] = []

    def fake_render_html_to_pdf(html_content: str) -> tuple[bytes, int]:
        preset_index = int(re.search(r'data-preset="(\d+)"', html_content).group(1))
        seen_presets.append(preset_index)
        page_counts = {0: 3, 1: 2, 2: 1}
        return (f"preset-{preset_index}".encode(), page_counts.get(preset_index, 1))

    monkeypatch.setattr(pdf_export, "_render_html_to_pdf", fake_render_html_to_pdf)

    pdf_bytes = pdf_export._generate_pdf_with_autofit_sync(
        "## Summary\nBuilt backend systems.\n",
        _personal_info(),
        "1_page",
    )

    assert pdf_bytes == b"preset-2"
    assert seen_presets[:3] == [0, 1, 2]
    assert all(index == 2 for index in seen_presets[3:])


def test_generate_pdf_autofit_prefers_tighter_density_before_smaller_font(monkeypatch):
    seen_presets: list[int] = []

    def fake_render_html_to_pdf(html_content: str) -> tuple[bytes, int]:
        preset_index = int(re.search(r'data-preset="(\d+)"', html_content).group(1))
        seen_presets.append(preset_index)
        page_counts = {0: 2, 1: 1, 2: 1}
        return (f"preset-{preset_index}".encode(), page_counts.get(preset_index, 1))

    monkeypatch.setattr(pdf_export, "_render_html_to_pdf", fake_render_html_to_pdf)

    pdf_bytes = pdf_export._generate_pdf_with_autofit_sync(
        "## Summary\nBuilt backend systems.\n",
        _personal_info(),
        "1_page",
    )

    assert pdf_bytes == b"preset-1"
    assert seen_presets[:2] == [0, 1]
    assert all(index == 1 for index in seen_presets[2:])


def test_generate_pdf_autofit_stops_at_minimum_preset(monkeypatch):
    seen_presets: list[int] = []

    def fake_render_html_to_pdf(html_content: str) -> tuple[bytes, int]:
        preset_index = int(re.search(r'data-preset="(\d+)"', html_content).group(1))
        seen_presets.append(preset_index)
        return (f"preset-{preset_index}".encode(), 5)

    monkeypatch.setattr(pdf_export, "_render_html_to_pdf", fake_render_html_to_pdf)

    pdf_bytes = pdf_export._generate_pdf_with_autofit_sync(
        "## Summary\nBuilt backend systems.\n",
        _personal_info(),
        "1_page",
    )

    min_preset_index = len(pdf_export.LAYOUT_PRESETS) - 1
    assert pdf_bytes == f"preset-{min_preset_index}".encode()
    assert seen_presets == list(range(len(pdf_export.LAYOUT_PRESETS)))


def test_generate_pdf_one_page_validation_tries_roomier_variant_before_export(monkeypatch):
    seen_fonts: list[float] = []

    monkeypatch.setattr(
        pdf_export,
        "LAYOUT_PRESETS",
        [pdf_export.LayoutPreset(body_font_size=10.0, line_height=1.10, page_margin=0.3, spacing_scale=0.7)],
    )

    def fake_render_html_to_pdf(html_content: str) -> tuple[bytes, int]:
        body_font_size = float(re.search(r"font-size:\s*([0-9.]+)pt;", html_content).group(1))
        seen_fonts.append(body_font_size)
        page_count = 1 if body_font_size <= 10.22 else 2
        return (f"font-{body_font_size:.2f}".encode(), page_count)

    monkeypatch.setattr(pdf_export, "_render_html_to_pdf", fake_render_html_to_pdf)

    pdf_bytes = pdf_export._generate_pdf_with_autofit_sync(
        "## Summary\nBuilt backend systems.\n",
        _personal_info(),
        "1_page",
    )

    assert pdf_bytes == b"font-10.22"
    assert seen_fonts[0] == 10.0
    assert 10.22 in seen_fonts
    assert 10.44 in seen_fonts
    assert len(seen_fonts) >= 3


def test_generate_pdf_one_page_validation_adds_section_spacing_when_room_exists(monkeypatch):
    seen_section_tops: list[float] = []
    accepted_section_tops: list[float] = []

    monkeypatch.setattr(
        pdf_export,
        "LAYOUT_PRESETS",
        [pdf_export.LayoutPreset(body_font_size=10.0, line_height=1.10, page_margin=0.3, spacing_scale=0.7)],
    )

    def fake_render_html_to_pdf(html_content: str) -> tuple[bytes, int]:
        section_margin_top = float(re.search(r"\.resume-section \{\s*margin-top: ([0-9.]+)pt;", html_content).group(1))
        body_font_size = float(re.search(r"font-size:\s*([0-9.]+)pt;", html_content).group(1))
        seen_section_tops.append(section_margin_top)
        # Allow section spacing growth, but force any font growth to overflow.
        page_count = 1 if body_font_size <= 10.0 else 2
        if page_count == 1:
            accepted_section_tops.append(section_margin_top)
        return (f"section-top-{section_margin_top:.2f}".encode(), page_count)

    monkeypatch.setattr(pdf_export, "_render_html_to_pdf", fake_render_html_to_pdf)

    pdf_bytes = pdf_export._generate_pdf_with_autofit_sync(
        "## Summary\nBuilt backend systems.\n",
        _personal_info(),
        "1_page",
    )

    assert seen_section_tops[0] > 0
    assert max(accepted_section_tops) > seen_section_tops[0]
    assert pdf_bytes == f"section-top-{max(accepted_section_tops):.2f}".encode()


def test_build_html_marks_dense_documents_in_html():
    html = pdf_export._build_html(
        (
            "## Experience\n"
            "- one\n- two\n- three\n- four\n- five\n"
            "## Projects\n"
            "- six\n- seven\n- eight\n- nine\n- ten\n"
            "## Skills\n"
            "- python\n- sql\n- aws\n- docker\n- linux\n"
        ),
        pdf_export.LAYOUT_PRESETS[0],
    )

    assert 'data-density="dense"' in html


def test_build_html_keeps_structured_entry_case_and_font_hierarchy():
    html = pdf_export._build_html(
        (
            "## Professional Experience\n"
            "Deloitte Canada | Toronto, Ontario, Canada\n"
            "Manager, Product Engineering | Jan 2022 - Present\n"
            "- Led platform work.\n"
        ),
        pdf_export.LAYOUT_PRESETS[0],
    )

    section_heading_size = pdf_export.LAYOUT_PRESETS[0].section_heading_size
    expected_primary_size = round(
        min(
            section_heading_size - 0.18,
            max(
                round(pdf_export.LAYOUT_PRESETS[0].body_font_size * 0.86, 2) + 0.32,
                pdf_export.LAYOUT_PRESETS[0].body_font_size * 0.90,
            ),
        ),
        2,
    )
    expected_content_size = round(pdf_export.LAYOUT_PRESETS[0].body_font_size * 0.86, 2)

    assert "DELOITTE CANADA" not in html
    assert "Deloitte Canada" in html
    assert expected_primary_size <= section_heading_size
    assert f"font-size: {expected_primary_size}pt;" in html
    assert f"font-size: {expected_content_size}pt;" in html


def test_build_html_bolds_only_professional_experience_role_title_split_rows():
    html = pdf_export._build_html(
        (
            "## Professional Experience\n"
            "Acme Corp | Toronto, ON\n"
            "Senior Data Architect | Jan 2020 - Present\n"
        ),
        pdf_export.LAYOUT_PRESETS[0],
    )

    assert "<span class='structured-left structured-left-primary'>Acme Corp</span>" in html
    assert "<span class='structured-left structured-left-secondary'>Senior Data Architect</span>" in html
    assert "Senior Data Architect" in html
    assert "Acme Corp" in html


def test_build_html_does_not_bold_date_split_rows_outside_professional_experience():
    html = pdf_export._build_html(
        (
            "## Summary\n"
            "Program Timeline | Jan 2020 - Present\n"
        ),
        pdf_export.LAYOUT_PRESETS[0],
    )

    assert "structured-entry structured-entry-summary" not in html


def test_build_html_preserves_blank_line_separation_between_split_row_groups():
    html = pdf_export._build_html(
        (
            "## Professional Experience\n"
            "Acme Corp | Toronto, ON\n"
            "Senior Data Architect | Jan 2020 - Present\n\n"
            "Beta Corp | Ottawa, ON\n"
            "Lead Engineer | Jan 2018 - Dec 2019\n"
        ),
        pdf_export.LAYOUT_PRESETS[0],
    )

    assert html.count("class='structured-entry structured-entry-professional_experience'") == 2


def test_build_export_document_reuses_one_normalized_header():
    document = pdf_export._build_export_document(
        "Alex Example\nalex@example.com | 555-0100\n\n## Summary\nBuilt backend systems.\n",
        _personal_info(),
    )

    assert document.header is not None
    assert document.header.name == "Alex Example"
    assert document.header.contact_line == "alex@example.com | 555-0100 | Toronto, ON | in/alex-example"
    assert len(document.sections) == 1
    assert document.sections[0].heading == "Summary"


def test_render_docx_sync_uses_letter_page_size_and_expected_margins():
    from docx import Document

    docx_bytes = pdf_export._render_docx_sync(
        "## Summary\nBuilt backend systems.\n",
        _personal_info(),
        "2_page",
    )
    document = Document(io.BytesIO(docx_bytes))
    section = document.sections[0]
    layout = pdf_export.DOCX_LAYOUT_PRESETS["2_page"]

    assert round(section.page_width.inches, 2) == 8.5
    assert round(section.page_height.inches, 2) == 11.0
    assert round(section.top_margin.inches, 2) == round(layout.page_margin, 2)
    assert round(section.left_margin.inches, 2) == round(layout.page_margin, 2)


def test_render_docx_sync_adds_clear_header_and_major_section_spacing():
    from docx import Document

    docx_bytes = pdf_export._render_docx_sync(
        (
            "# Alex Example\n"
            "alex@example.com | 555-0100 | Toronto, ON\n\n"
            "## Summary\n"
            "Built backend systems.\n\n"
            "## Skills\n"
            "- Python\n"
        ),
        _personal_info(),
        "2_page",
    )
    document = Document(io.BytesIO(docx_bytes))
    layout = pdf_export.DOCX_LAYOUT_PRESETS["2_page"]
    header_section_gap = max(layout.body_font_size * layout.line_spacing * 0.95, layout.header_spacing_after)
    major_section_gap = max(layout.body_font_size * layout.line_spacing * 0.95, layout.section_spacing_before)
    contact_paragraph = next(paragraph for paragraph in document.paragraphs if "alex@example.com" in paragraph.text)
    skills_heading = next(paragraph for paragraph in document.paragraphs if paragraph.text == "SKILLS")

    assert round(contact_paragraph.paragraph_format.space_after.pt, 2) >= round(header_section_gap, 2)
    assert round(skills_heading.paragraph_format.space_before.pt, 2) >= round(major_section_gap, 2)


def test_render_docx_sync_renders_header_bullets_and_split_rows_without_tables():
    from docx import Document

    docx_bytes = pdf_export._render_docx_sync(
        (
            "# Alex Example\n"
            "alex@example.com | 555-0100 | Toronto, ON\n\n"
            "## Professional Experience\n"
            "Acme Corp | Toronto, ON\n"
            "Senior Data Architect | Jan 2020 - Present\n"
            "- Led migration program\n"
        ),
        _personal_info(),
        "1_page",
    )
    document = Document(io.BytesIO(docx_bytes))
    paragraph_text = [paragraph.text for paragraph in document.paragraphs]

    assert document.tables == []
    assert paragraph_text[0] == "Alex Example"
    assert any("PROFESSIONAL EXPERIENCE" == text for text in paragraph_text)
    assert any("Acme Corp\tToronto, ON" == text for text in paragraph_text)
    assert any("Senior Data Architect\tJan 2020 - Present" == text for text in paragraph_text)
    assert any("Led migration program" in text for text in paragraph_text)


def test_render_docx_sync_bolds_only_professional_experience_role_title_split_rows():
    from docx import Document

    docx_bytes = pdf_export._render_docx_sync(
        (
            "## Professional Experience\n"
            "Acme Corp | Toronto, ON\n"
            "Senior Data Architect | Jan 2020 - Present\n"
            "## Summary\n"
            "Portfolio Lead | 2024 - Present\n"
        ),
        _personal_info(),
        "1_page",
    )
    document = Document(io.BytesIO(docx_bytes))
    company_paragraph = next(paragraph for paragraph in document.paragraphs if paragraph.text == "Acme Corp\tToronto, ON")
    role_paragraph = next(paragraph for paragraph in document.paragraphs if paragraph.text == "Senior Data Architect\tJan 2020 - Present")
    summary_paragraph = next(paragraph for paragraph in document.paragraphs if paragraph.text == "Portfolio Lead\t2024 - Present")

    assert company_paragraph.runs[0].bold is True
    assert role_paragraph.runs[0].italic is True
    assert summary_paragraph.runs[0].bold in {False, None}


def test_resolve_docx_layout_adjusts_for_dense_and_sparse_documents():
    dense = pdf_export._resolve_docx_layout("1_page", {"is_dense": True, "is_sparse": False})
    sparse = pdf_export._resolve_docx_layout("1_page", {"is_dense": False, "is_sparse": True})
    baseline = pdf_export.DOCX_LAYOUT_PRESETS["1_page"]

    assert dense.body_font_size < baseline.body_font_size
    assert dense.paragraph_spacing < baseline.paragraph_spacing
    assert sparse.body_font_size > baseline.body_font_size
    assert sparse.section_spacing_before > baseline.section_spacing_before
